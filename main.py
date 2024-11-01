import streamlit as st
import json
import redis
from utils.pdf_processing import process_pdf_pages, process_pdf_task
from utils.llm_interaction import ask_question
from concurrent.futures import ThreadPoolExecutor, as_completed
import asyncio
import io
import tiktoken
from docx import Document

# Initialize Redis client
redis_client = redis.StrictRedis(
    host="yuktestredis.redis.cache.windows.net",  # Replace with your Redis host
    port=6380,  # Redis SSL port
    password="VBhswgzkLiRpsHVUf4XEI2uGmidT94VhuAzCaB2tVjs=",  # Replace with your Redis password
    ssl=False,  # Enable SSL for Azure Redis
)

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []


def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    return len(tokens)


async def handle_question(prompt, spinner_placeholder):
    if prompt:
        try:
            input_tokens = count_tokens(prompt)
            document_data = retrieve_document_data()  # Retrieve data from Redis
            document_tokens = count_tokens(json.dumps(document_data))
            total_input_tokens = input_tokens + document_tokens

            with spinner_placeholder.container():
                st.markdown(
                    """
                    <header>
                    <div style="text-align: center;">
                        <div class="spinner" style="margin: 20px;">
                            <div class="bounce1"></div>
                            <div class="bounce2"></div>
                            <div class="bounce3"></div>
                        </div>
                    </div>
                    </header>
                    """,
                    unsafe_allow_html=True,
                )

                answer = await asyncio.get_event_loop().run_in_executor(
                    None,
                    ask_question,
                    document_data,
                    prompt,
                    st.session_state.chat_history,
                )

            output_tokens = count_tokens(answer)
            st.session_state.chat_history.append(
                {
                    "question": prompt,
                    "answer": answer,
                    "input_tokens": total_input_tokens,
                    "output_tokens": output_tokens,
                }
            )

        except Exception as e:
            st.error(f"Error processing question: {e}")
        finally:
            spinner_placeholder.empty()


def store_document_data(doc_name, data):
    """Store processed document data in Redis."""
    redis_client.set(doc_name, json.dumps(data))


def retrieve_document_data():
    """Retrieve all document data from Redis."""
    document_data = {}
    for doc_name in redis_client.keys():
        document_data[doc_name.decode("utf-8")] = json.loads(redis_client.get(doc_name))
    return document_data


def reset_session():
    st.session_state.chat_history = []
    st.session_state.uploaded_files = []
    redis_client.flushdb()  # Clear all documents from Redis


def display_chat():
    if st.session_state.chat_history:
        for i, chat in enumerate(st.session_state.chat_history):
            user_message = f"""
            <div style='padding:10px; border-radius:10px; margin:5px 0; text-align:right;'>
            {chat['question']}
            <small style='color:grey;'>Tokens: {chat['input_tokens']}</small></div>
            """
            assistant_message = f"""
            <div style='padding:10px; border-radius:10px; margin:5px 0; text-align:left;'>
            {chat['answer']}
            <small style='color:grey;'>Tokens: {chat['output_tokens']}</small></div>
            """
            st.markdown(user_message, unsafe_allow_html=True)
            st.markdown(assistant_message, unsafe_allow_html=True)

            chat_content = {
                "question": chat["question"],
                "answer": chat["answer"],
                "input_tokens": chat["input_tokens"],
                "output_tokens": chat["output_tokens"],
            }
            doc = generate_word_document(chat_content)
            word_io = io.BytesIO()
            doc.save(word_io)
            word_io.seek(0)

            st.download_button(
                label="↴",
                data=word_io,
                file_name=f"chat_{i+1}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


def generate_word_document(content):
    doc = Document()
    doc.add_heading("Chat Response", 0)
    doc.add_paragraph(f"Question: {content['question']}")
    doc.add_paragraph(f"Answer: {content['answer']}")
    doc.add_paragraph(f"Input Tokens: {content['input_tokens']}")
    doc.add_paragraph(f"Output Tokens: {content['output_tokens']}")
    return doc


with st.sidebar:
    uploaded_files = st.file_uploader(
        "Upload your documents",
        type=["pdf", "docx", "xlsx", "pptx"],
        accept_multiple_files=True,
        help="Supports PDF, DOCX, XLSX, and PPTX formats.",
    )
    if uploaded_files:
        new_files = []
        for index, uploaded_file in enumerate(uploaded_files):
            if not redis_client.exists(
                uploaded_file.name
            ):  # Check Redis instead of session state
                new_files.append(uploaded_file)
            else:
                st.info(f"{uploaded_file.name} is already uploaded.")

        if new_files:
            progress_text = st.empty()
            progress_bar = st.progress(0)
            total_files = len(new_files)

            with st.spinner("Learning about your document(s)..."):
                with ThreadPoolExecutor(max_workers=2) as executor:
                    future_to_file = {
                        executor.submit(
                            process_pdf_task, uploaded_file, first_file=(index == 0)
                        ): uploaded_file
                        for index, uploaded_file in enumerate(new_files)
                    }

                    for i, future in enumerate(as_completed(future_to_file)):
                        uploaded_file = future_to_file[future]
                        try:
                            document_data = future.result()
                            store_document_data(
                                uploaded_file.name, document_data
                            )  # Store data in Redis
                            st.success(f"{uploaded_file.name} processed successfully!")
                        except Exception as e:
                            st.error(f"Error processing {uploaded_file.name}: {e}")

                        progress_bar.progress((i + 1) / total_files)

            progress_text.text("Processing complete.")
            progress_bar.empty()

    if redis_client.keys():  # Check Redis for existing documents
        download_data = json.dumps(retrieve_document_data(), indent=4)
        st.download_button(
            label="Download Document Analysis",
            data=download_data,
            file_name="document_analysis.json",
            mime="application/json",
        )

st.image("logoD.png", width=200)
st.title("docQuest")
st.subheader("Unveil the Essence, Compare Easily, Analyze Smartly", divider="orange")

if redis_client.keys():
    prompt = st.chat_input("Ask me anything about your documents", key="chat_input")
    spinner_placeholder = st.empty()
    if prompt:
        asyncio.run(handle_question(prompt, spinner_placeholder))

display_chat()

total_input_tokens = sum(chat["input_tokens"] for chat in st.session_state.chat_history)
total_output_tokens = sum(
    chat["output_tokens"] for chat in st.session_state.chat_history
)
st.sidebar.write(f"Total Input Tokens: {total_input_tokens}")
st.sidebar.write(f"Total Output Tokens: {total_output_tokens}")
